"""
One-pass JAMIA submission formatter for AToP_paper_v6.docx.

What this does (in order):
  1. Loads the 218,196-patient one-per-patient cache
     (../data/mimiciv/3.1/hosp/.atop_cache/samples_*.pkl with label sum = 27,508)
     and computes Table 1 stats directly from it — no manual filling.
  2. Copies figures to ../manuscript/jamia_submission/figures/ with
     JAMIA-standard names (Figure1.pptx, Figure2/3/4.png).
  3. Fills Table 1, Table 2, Table 3 [val] cells positionally
     (the docx uses generic [val], not named placeholders).
  4. Removes the Acuity (emergency) row from Table 1 (redundant with
     Emergency/Urgent).
  5. Fixes XGB PR-AUC 0.32 -> 0.31 everywhere.
  6. Saves AToP_JAMIA_submission.docx.

Text edits B-G and the Figure Legends section are applied by the
companion script scripts/jamia_final_edits.py (see TASKS/jamia_final_edits.md)
so that the number-audit pass can also catch any stale prose values
(e.g. "2%" → "0.1%" for the mining support threshold).

Run from atop_package/:
    python scripts/prepare_jamia_submission.py
"""
from __future__ import annotations
import os, re, sys, glob, pickle, shutil
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# ── Paths ────────────────────────────────────────────────────────────────
SRC_CANDIDATES = [
    "../manuscript/atop_v5.5/AToP_paper_v6.docx",
    "../manuscript/claude files/AToP_paper_v6.docx",
]
SRC = next((p for p in SRC_CANDIDATES if os.path.exists(p)), None)
if SRC is None:
    raise FileNotFoundError(
        f"AToP_paper_v6.docx not found in any of: {SRC_CANDIDATES}"
    )
SUBMIT_DIR = "../manuscript/jamia_submission"
FIG_DIR    = f"{SUBMIT_DIR}/figures"
DST        = f"{SUBMIT_DIR}/AToP_JAMIA_submission.docx"
os.makedirs(FIG_DIR, exist_ok=True)

FULL_RUN   = "../runs/full_CPD"
MIMIC_DIR  = "../data/mimiciv/3.1/hosp"
PKL_CACHE  = f"{MIMIC_DIR}/.atop_cache"

# ── Helpers ─────────────────────────────────────────────────────────────
def set_cell(cell, text):
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.add_run(text)
    for extra in cell.paragraphs[1:]:
        if extra.runs:
            extra.runs[0].text = ""
            for r in extra.runs[1:]:
                r.text = ""

def replace_in_para(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True

def replace_all(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        if replace_in_para(p, old, new):
            n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_para(p, old, new):
                        n += 1
    return n

def append_to_para(para, suffix):
    if para.runs:
        full = "".join(r.text for r in para.runs) + suffix
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


# ────────────────────────────────────────────────────────────────────────
# Step A — Load the 218,196-patient pkl and compute Table 1 stats
# ────────────────────────────────────────────────────────────────────────
print("=" * 70)
print("[A] Loading one-per-patient pkl and computing Table 1 stats")
print("=" * 70)

TARGET_N = 218196  # matches summary.json
pkls = sorted(glob.glob(f"{PKL_CACHE}/samples_*.pkl"))
pkl_path = None
for p in pkls:
    with open(p, "rb") as f:
        d = pickle.load(f)
    if len(d) == TARGET_N:
        pkl_path = p
        samples = d
        break
if pkl_path is None:
    raise FileNotFoundError(
        f"No samples_*.pkl of size {TARGET_N} found in {PKL_CACHE}"
    )
print(f"    pkl: {os.path.basename(pkl_path)}  ({len(samples):,} samples)")

labels    = np.array([s["readmit_30d"] for s in samples], dtype=float)
n_visits  = np.array([s.get("n_visits", 0) for s in samples], dtype=int)
pids      = [str(s["patient_id"]) for s in samples]
hadm_ids  = [str(s["index_hadm_id"]) for s in samples]
cohort = pd.DataFrame({
    "subject_id": pids,
    "hadm_id": hadm_ids,
    "readmit_30d": labels,
    "n_visits": n_visits,
})
n_total = len(cohort)
n_r     = int(cohort["readmit_30d"].sum())
n_nr    = n_total - n_r
print(f"    total={n_total:,}  readmit={n_r:,}  no_readmit={n_nr:,}")
assert n_total == TARGET_N, f"pkl size ({n_total}) != TARGET_N ({TARGET_N})"

# ── Merge demographics + admission type + LACE ────────────────────────────
patients = pd.read_csv(
    f"{MIMIC_DIR}/patients.csv.gz",
    dtype={"subject_id": str},
    usecols=["subject_id", "gender", "anchor_age"],
)
admissions = pd.read_csv(
    f"{MIMIC_DIR}/admissions.csv.gz",
    dtype={"subject_id": str, "hadm_id": str},
    usecols=["subject_id", "hadm_id", "admission_type"],
)
lace = pd.read_csv(
    f"{FULL_RUN}/lace_scores_all.csv",
    dtype={"subject_id": str, "hadm_id": str},
)

cohort = cohort.merge(patients, on="subject_id", how="left")
cohort = cohort.merge(admissions, on=["subject_id", "hadm_id"], how="left")
cohort = cohort.merge(
    lace[["subject_id", "hadm_id", "los_days", "cci", "ed_visits_6mo", "lace_score"]],
    on=["subject_id", "hadm_id"], how="left",
)
cohort["age"]    = pd.to_numeric(cohort["anchor_age"], errors="coerce")
cohort["female"] = (cohort["gender"] == "F").astype(float)
elective_types = {"ELECTIVE", "SURGICAL SAME DAY ADMISSION"}
cohort["is_emergency"] = (
    ~cohort["admission_type"].str.upper().isin(elective_types)
).astype(float)

r  = cohort[cohort["readmit_30d"] == 1]
nr = cohort[cohort["readmit_30d"] == 0]

def _med_iqr_int(s):
    s = pd.to_numeric(s, errors="coerce").dropna()
    return f"{s.median():.0f} ({s.quantile(0.25):.0f}–{s.quantile(0.75):.0f})"

def _med_iqr_float(s):
    s = pd.to_numeric(s, errors="coerce").dropna()
    return f"{s.median():.1f} ({s.quantile(0.25):.1f}–{s.quantile(0.75):.1f})"

def _n_pct(s, n_cohort):
    s = pd.to_numeric(s, errors="coerce").dropna()
    return f"{int(s.sum()):,} ({s.sum()/n_cohort*100:.1f}%)"

T1 = {
    "age":        (_med_iqr_int(cohort["age"]),
                   _med_iqr_int(r["age"]),
                   _med_iqr_int(nr["age"])),
    "female":     (_n_pct(cohort["female"], len(cohort)),
                   _n_pct(r["female"], len(r)),
                   _n_pct(nr["female"], len(nr))),
    "emergency":  (_n_pct(cohort["is_emergency"], len(cohort)),
                   _n_pct(r["is_emergency"], len(r)),
                   _n_pct(nr["is_emergency"], len(nr))),
    "los":        (_med_iqr_float(cohort["los_days"]),
                   _med_iqr_float(r["los_days"]),
                   _med_iqr_float(nr["los_days"])),
    "cci":        (_med_iqr_int(cohort["cci"]),
                   _med_iqr_int(r["cci"]),
                   _med_iqr_int(nr["cci"])),
    "ed":         (_med_iqr_int(cohort["ed_visits_6mo"]),
                   _med_iqr_int(r["ed_visits_6mo"]),
                   _med_iqr_int(nr["ed_visits_6mo"])),
    "prior_adm":  (_med_iqr_int(cohort["n_visits"]),
                   _med_iqr_int(r["n_visits"]),
                   _med_iqr_int(nr["n_visits"])),
    "lace":       (_med_iqr_int(cohort["lace_score"]),
                   _med_iqr_int(r["lace_score"]),
                   _med_iqr_int(nr["lace_score"])),
}
for k, v in T1.items():
    print(f"    {k:<10} overall={v[0]}  readmit={v[1]}  no_readmit={v[2]}")


# ────────────────────────────────────────────────────────────────────────
# Step B — Copy figures
# ────────────────────────────────────────────────────────────────────────
print("\n" + "=" * 70)
print("[B] Copying figures")
print("=" * 70)
FIG_SRC_DIR = os.path.dirname(SRC)
FIG_MAIN    = f"{FULL_RUN}/explain/figures/main"
fig_map = [
    (f"{FIG_SRC_DIR}/atop_framework_fig1.pptx",           "Figure1.pptx"),
    (f"{FIG_SRC_DIR}/atop_framework_fig1.html",           "Figure1_interactive.html"),
    (f"{FIG_MAIN}/fig2_patient_ig_tokens_test.png",       "Figure2.png"),
    (f"{FIG_MAIN}/fig3_main_test_j20_short.png",          "Figure3.png"),
    (f"{FIG_MAIN}/fig4_dumbbell_test_j20_short.png",      "Figure4.png"),
]
for src, dst_name in fig_map:
    dst_path = os.path.join(FIG_DIR, dst_name)
    if os.path.exists(src):
        shutil.copyfile(src, dst_path)
        print(f"    ✅ {dst_name}  ({os.path.getsize(dst_path)/1024:.0f} KB)")
    else:
        print(f"    ⚠️  missing: {src}")


# ────────────────────────────────────────────────────────────────────────
# Step C — Open doc, fill Table 1/2/3, run text edits, add legends
# ────────────────────────────────────────────────────────────────────────
print("\n" + "=" * 70)
print("[C] Opening doc and applying edits")
print("=" * 70)
doc = Document(SRC)

# ── C1. Fix XGB PR-AUC 0.32 → 0.31 ──
print("\n[C1] XGB PR-AUC 0.32 → 0.31")
for old, new in [
    ("PR-AUC 0.32",    "PR-AUC 0.31"),
    ("PR-AUC ~0.32",   "PR-AUC ~0.31"),
    ("PR-AUC of 0.32", "PR-AUC of 0.31"),
    ("(PR-AUC 0.32",   "(PR-AUC 0.31"),
    ("PR-AUC: 0.32",   "PR-AUC: 0.31"),
]:
    n = replace_all(doc, old, new)
    if n:
        print(f"    Fixed '{old}' ({n}x)")

# ── C2. Table 1 (cohort characteristics) — full-cohort values ──
print("\n[C2] Table 1 — filling with full 218,196 cohort")
t0 = doc.tables[0]
HEADER = {
    "Readmitted":     f"Readmitted\n(n={n_r:,})",
    "Not Readmitted": f"Not Readmitted\n(n={n_nr:,})",
}
for cell in t0.rows[0].cells:
    text = cell.text.strip()
    if "[val]" in text:
        for label, new_text in HEADER.items():
            if text.startswith(label):
                set_cell(cell, new_text)
                break

# row_idx → (Overall, Readmitted, Not Readmitted, P)
t0_rows = {
    2:  (*T1["age"],       "—"),   # Age
    3:  (*T1["female"],    "—"),   # Female sex
    5:  (*T1["emergency"], "—"),   # Emergency/Urgent
    7:  (*T1["los"],       "—"),   # LOS (days)
    8:  "DELETE",                  # Acuity — delete row (redundant)
    9:  (*T1["cci"],       "—"),   # CCI
    10: (*T1["ed"],        "—"),   # ED visits 6mo
    11: (*T1["prior_adm"], "—"),   # Prior admissions
}

# Apply row updates first (leave Acuity for deletion afterward)
for ri, values in t0_rows.items():
    if values == "DELETE":
        continue
    if ri >= len(t0.rows):
        continue
    row = t0.rows[ri]
    for ci, val in enumerate(values, start=1):
        if ci < len(row.cells):
            set_cell(row.cells[ci], val)
print(f"    Header + rows 2/3/5/7/9/10/11 filled")

# Delete the Acuity row by unique content match (row-index may shift)
deleted_acuity = False
for row in list(t0.rows):
    row_text = " ".join(c.text.lower() for c in row.cells)
    if "acuity" in row_text:
        t0._tbl.remove(row._tr)
        deleted_acuity = True
        print(f"    Deleted Acuity row (redundant with Emergency/Urgent)")
        break
if not deleted_acuity:
    print("    ⚠️  Acuity row not found")


# ── C3. Table 1 (Mining summary) ──
print("\n[C3] Table 2 — mining summary")
t1 = doc.tables[1]
t1_rows = {
    2: ("36.1 (43.3)",  "—"),
    3: ("13.7 (12.4)",  "—"),
    4: ("38.0",         "—"),
    6: ("24,490",       "—"),
    7: ("5,325",        "—"),
    8: ("5,047",        "—"),
    10: ("—", "1,877"),
    11: ("—", "3,170"),
}
for ri, (train, test) in t1_rows.items():
    if ri >= len(t1.rows):
        continue
    row = t1.rows[ri]
    if len(row.cells) >= 3:
        set_cell(row.cells[1], train)
        set_cell(row.cells[2], test)
print("    All 8 mining metrics filled")


# ── C4. Table 2 (Phenotype clusters) — ΣIG, n carriers, Δŷ ──
print("\n[C4] Table 3 — phenotype clusters")
dumbbell = pd.read_csv(f"{FIG_MAIN}/fig4_dumbbell_test_j20_short.csv")
pathway  = pd.read_csv(f"{FULL_RUN}/explain/figures/csv/fig3_pathway_importance_test.csv")

def _canon(s: str) -> str:
    t = str(s).lower().strip()
    t = t.replace("→", "->").replace("–", "-").replace("—", "-")
    t = re.sub(r"\s*\([^)]*\)", "", t)
    t = re.sub(r"[\{\},\s\.]", "", t)
    return t

label_to_data = {}
for _, row in dumbbell.iterrows():
    pat = row["pattern"]
    lbl = row["label"]
    sig = float(row["original_sig_ig"]) if pd.notna(row.get("original_sig_ig")) else None
    dlt = float(row["delta"])           if pd.notna(row.get("delta"))           else None
    n = None
    m = pathway[pathway["pattern"] == pat]
    if len(m):
        n = int(m["n_present"].iloc[0])
    label_to_data[_canon(lbl)] = (sig, n, dlt)

ALIASES = {
    _canon("{Live birth, Dibucaine} → Live birth"):
        _canon("{Live birth, RX_3339} → Live birth"),
    _canon("{Chest pain, Resp meas.} → Chest pain"):
        _canon("{Chest pain, Resp measurements} → Chest pain"),
    _canon("{MDD, SI} → SI"):
        _canon("{MDD (single), SI} → SI"),
}
for a, b in ALIASES.items():
    if b in label_to_data and a not in label_to_data:
        label_to_data[a] = label_to_data[b]

t2 = doc.tables[2]
filled = 0
for ri in range(1, len(t2.rows)):
    row = t2.rows[ri]
    if len(row.cells) < 5:
        continue
    pat_text = row.cells[1].text.strip()
    if not pat_text:
        continue
    key = _canon(pat_text)
    sig = n = dlt = None
    if key in label_to_data:
        sig, n, dlt = label_to_data[key]
    else:
        for k2, (s2, n2, d2) in label_to_data.items():
            if key in k2 or k2 in key:
                sig, n, dlt = s2, n2, d2
                break
    if sig is not None and "[val]" in row.cells[2].text:
        set_cell(row.cells[2], f"{sig:+.4f}")
        filled += 1
    if n is not None and "[val]" in row.cells[3].text:
        set_cell(row.cells[3], f"{n}")
        filled += 1
    if dlt is not None and len(row.cells) > 4 and "[val]" in row.cells[4].text:
        set_cell(row.cells[4], f"{dlt:+.3f}")
        filled += 1
print(f"    Filled {filled} cells")


# ── C5. Caption note replacement ──
for old, new in [
    ("[val] = to be filled from pipeline CSVs.",
     "(Values sourced from runs/full_CPD/explain/ cached CSVs.)"),
    ("[val] = to be filled from pipeline.",
     "(Values sourced from runs/full_CPD/explain/ cached CSVs.)"),
]:
    replace_all(doc, old, new)


# ── C8. Remaining placeholder scan ──
print("\n[C8] Remaining placeholder scan")
IGNORE = {"[VISIT]", "[CLS]", "[PAD]", "[IG]", "[CPD]"}
def _placeholders(text):
    raw = re.findall(r"\[val\]|\[[A-Z_]+\]", text)
    return [m for m in raw if m not in IGNORE]

remaining = []
for i, para in enumerate(doc.paragraphs):
    hits = _placeholders(para.text)
    if hits:
        remaining.append((f"p{i}", para.text[:120], hits))
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            hits = _placeholders(cell.text)
            if hits:
                remaining.append((f"T{ti}r{ri}c{ci}", cell.text[:120], hits))

if remaining:
    print(f"    ⚠️  {len(remaining)} placeholder(s) remain:")
    for loc, snip, hits in remaining:
        print(f"      [{loc}] {hits}  |  '{snip}'")
else:
    print("    ✅ None remaining")


# ── C9. Abstract word count ──
print("\n[C9] Abstract word count")
MAJOR = ("background and significance", "introduction", "keywords")
in_abs, count = False, 0
for para in doc.paragraphs:
    t = para.text.strip()
    if not in_abs and t.lower() == "abstract":
        in_abs = True
        continue
    if in_abs:
        if any(t.lower().startswith(k) for k in MAJOR):
            break
        count += len(t.split())
print(f"    ~{count} words  " + ("✅ ≤250" if count <= 250 else f"⚠️ over by {count-250}"))


# ── D. Save ──
print("\n" + "=" * 70)
print(f"[D] Saving → {DST}")
print("=" * 70)
doc.save(DST)
print(f"    ✅ Saved ({os.path.getsize(DST)/1024:.0f} KB)")

print("\n" + "=" * 70)
print("TABLE 1 — numbers written to docx (copy into any external text):")
print("=" * 70)
print(f"  N total:       {n_total:,}")
print(f"  N readmitted:  {n_r:,}")
print(f"  N not readmit: {n_nr:,}")
for row_name, vals in [
    ("Age",           T1["age"]),
    ("Female",        T1["female"]),
    ("Emergency/Urg", T1["emergency"]),
    ("LOS (days)",    T1["los"]),
    ("CCI",           T1["cci"]),
    ("ED visits 6mo", T1["ed"]),
    ("Prior adm",     T1["prior_adm"]),
    ("LACE score",    T1["lace"]),
]:
    print(f"  {row_name:<14} overall={vals[0]:<18}  readmit={vals[1]:<18}  no_readm={vals[2]}")
