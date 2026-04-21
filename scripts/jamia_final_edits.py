"""
Final edits to AToP_JAMIA_submission.docx:
  1. Audit and fix stale numbers in prose (partial-cohort values, 2% → 0.1%)
  2. Apply text edits B–G
  3. Add Figure Legends section after References

Run from atop_package/, after scripts/prepare_jamia_submission.py.
"""
from __future__ import annotations
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

SRC = "../manuscript/jamia_submission/AToP_JAMIA_submission.docx"
DST = "../manuscript/jamia_submission/AToP_JAMIA_submission.docx"  # overwrite

doc = Document(SRC)


def replace_in_para(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_all(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        if replace_in_para(p, old, new):
            n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_para(p, old, new):
                        n += 1
    return n


# ── 1. Number audit — fix stale values in prose ───────────────────────────────
print("[1] Auditing stale numbers in prose...")
NUMBER_FIXES = [
    # Readmitted / not readmitted counts
    ("24,876",         "27,508"),
    ("171,501",        "190,688"),
    ("196,377",        "218,196"),
    # Rates and medians that changed
    ("78.4%",          "89.6%"),
    ("4.0 (2.0–8.0)",  "2.6 (1.0–5.1)"),
    ("6.0 (3.0–11.0)", "3.2 (1.2–6.8)"),
    ("4.0 (2.0–7.0)",  "2.6 (1.0–4.9)"),
    ("2.0 (0.0–5.0)",  "1 (0–2)"),
    ("4.0 (1.0–7.0)",  "1 (0–4)"),
    ("2.0 (0.0–4.0)",  "1 (0–2)"),
    # Female counts that shifted
    ("12,489 (50.2%)", "14,141 (51.4%)"),
    ("91,370 (53.3%)", "101,203 (53.1%)"),
    ("115,461 (52.9%)", "115,344 (52.9%)"),
    # Mining threshold — config.json said 0.02 but actual run used 0.001
    ("2% minimum support",       "0.1% minimum support"),
    ("minimum support of 2%",    "minimum support of 0.1%"),
    ("support threshold of 2%",  "support threshold of 0.1%"),
    ("min_support_frac=0.02",    "min_support_frac=0.001"),
    ("min_support_frac = 0.02",  "min_support_frac = 0.001"),
]
for old, new in NUMBER_FIXES:
    n = replace_all(doc, old, new)
    if n:
        print(f"    Fixed: '{old}' → '{new}'  ({n}x)")

# ── 2. Edit B — Mining methods ────────────────────────────────────────────────
# Insert the case/control-separated mining description after the anchor
# "with minimum support filtering" in the Methods-body pipeline paragraph
# (the abstract's mining summary uses different wording so won't match).
print("\n[2] Edit B: Mining methods clarification...")
B_APPEND = (
    " Episode mining was applied separately to readmitted and non-readmitted "
    "training patients, each with a minimum support threshold of 0.1% of its "
    "group (≥20 cases, ≥134 controls); discovered patterns were pooled and "
    "cross-verified across groups, retaining only patterns carried by ≥0.1% "
    "of the full training set (n≥153 of 152,737)."
)
added_b = False
for para in doc.paragraphs:
    if "with minimum support filtering" in para.text:
        full = "".join(r.text for r in para.runs)
        anchor = "with minimum support filtering"
        i = full.find(anchor) + len(anchor)
        new_full = full[:i] + B_APPEND + full[i:]
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""
        added_b = True
        print("    Appended mining clarification after 'minimum support filtering'")
        break
if not added_b:
    print("    ⚠️  Anchor 'with minimum support filtering' not found")
    print(f"    Text to paste: {B_APPEND.strip()}")

# ── 3. Edit C — LACE+ description ─────────────────────────────────────────────
# Use the unique substring that only appears in the body Results paragraph.
print("\n[3] Edit C: LACE+ reframe...")
C_OLD = (
    "outperforming the LACE+ index [7], a widely used readmission risk score "
    "incorporating length of stay, acuity, comorbidity, ED utilization, age, "
    "sex, and prior admissions: LACE+ logistic regression"
)
C_NEW = (
    "outperforming logistic regression and XGBoost classifiers trained on "
    "LACE-derived features (length of stay, acuity, Charlson Comorbidity "
    "Index, ED visits in the prior 6 months, age, sex, and prior admission "
    "count): LACE+ logistic regression"
)
if replace_all(doc, C_OLD, C_NEW):
    print("    Reframed LACE+ as ML classifiers trained on LACE-derived features")
else:
    print("    ⚠️  LACE+ sentence not matched — add manually")
    print(f"    Text to paste: {C_NEW}")

# ── 4. Edit D — Index admission caveat ───────────────────────────────────────
print("\n[4] Edit D: Index admission caveat...")
D_CAVEAT = (
    " This ensures an observed future window for readmission ascertainment "
    "within the dataset, though it may under-represent end-of-trajectory "
    "clinical patterns."
)
anchor_d = "with all prior admissions serving as longitudinal history."
added_d = False
for para in doc.paragraphs:
    if anchor_d in para.text:
        full = "".join(r.text for r in para.runs)
        new_full = full.replace(anchor_d, anchor_d + D_CAVEAT)
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""
        added_d = True
        print("    Added caveat after 'longitudinal history.'")
        break
if not added_d:
    print("    ⚠️  Penultimate-admission anchor not found")
    print(f"    Text to paste: {D_CAVEAT.strip()}")

# ── 5. Edit E — Table 3 footnote ─────────────────────────────────────────────
print("\n[5] Edit E: Table 3 footnote...")
T3_FOOTNOTE = "Pattern ranking on the test set is based entirely on label-free ΣIG attribution."
added = False
for i, tbl in enumerate(doc.tables):
    header = " ".join(c.text.lower() for c in tbl.rows[0].cells) if tbl.rows else ""
    if any(kw in header for kw in ["pattern", "sig_ig", "sig ig", "carrier", "delta"]):
        tbl_elem = tbl._tbl
        parent   = tbl_elem.getparent()
        idx      = list(parent).index(tbl_elem)
        p_elem   = OxmlElement("w:p")
        r_elem   = OxmlElement("w:r")
        t_elem   = OxmlElement("w:t")
        t_elem.text = T3_FOOTNOTE
        r_elem.append(t_elem)
        p_elem.append(r_elem)
        parent.insert(idx + 1, p_elem)
        added = True
        print(f"    Added Table 3 footnote after table {i}")
        break
if not added:
    print("    ⚠️  Table 3 not identified — add footnote manually")
    print(f"    Text to paste: {T3_FOOTNOTE}")

# ── 6. Edit F — Masking caveat (target ONLY the "Token masking:" subsection) ─
print("\n[6] Edit F: Masking caveat...")
F_CAVEAT = (
    " Masking changes inputs outside the training distribution; results are "
    "interpreted as evidence of model sensitivity rather than causal reliance."
)
added_f = False
for para in doc.paragraphs:
    if para.text.startswith("Token masking:"):
        full = "".join(r.text for r in para.runs) + F_CAVEAT
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ""
        added_f = True
        print("    Appended caveat to 'Token masking:' paragraph")
        break
if not added_f:
    print("    ⚠️  'Token masking:' paragraph not found")
    print(f"    Text to paste: {F_CAVEAT.strip()}")

# ── 7. Edit G — Figure 3 caption panel descriptions ──────────────────────────
# The caption paragraph starts with "[Figure 3:" in square brackets — unique.
print("\n[7] Edit G: Figure 3 caption...")
G_APPEND = (
    " Panel A: which tokens matter at the population level (prevalence-weighted). "
    "Panel B: which tokens have the strongest attribution when present (conditional signal). "
    "Panel C: which temporal configurations concentrate attribution among carriers "
    "(structured narratives)."
)
added_g = False
for para in doc.paragraphs:
    if para.text.startswith("[Figure 3:") or para.text.startswith("[Figure 3."):
        full = "".join(r.text for r in para.runs) + G_APPEND
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ""
        added_g = True
        print("    Appended panel descriptions to Figure 3 caption")
        break
if not added_g:
    print("    ⚠️  Figure 3 caption not found")
    print(f"    Text to paste: {G_APPEND.strip()}")

# ── 8. Add Figure Legends section after References ────────────────────────────
print("\n[8] Adding Figure Legends section...")

LEGENDS = [
    ("Figure Legends",
     True, False, 13, 20),

    ("Figure 1. AToP framework overview.",
     True, False, 11, 12),
    ("Four-stage pipeline illustrating the AToP framework: (1) Attribution — "
     "per-token Integrated Gradients computed on the trained Transformer; "
     "(2) Mapping — salient tokens mapped back to their admission-block structure; "
     "(3) Mining — frequent cross-visit temporal patterns extracted from training-set "
     "salient blocks; (4) Validation — patterns evaluated on held-out test data via "
     "token masking.",
     False, False, 11, 2),
    ("Alt text: Horizontal four-stage diagram showing the AToP pipeline — "
     "Attribution, Mapping, Mining, and Validation — connected by left-to-right arrows.",
     False, True, 10, 2),

    ("Figure 2. Single-patient attribution example.",
     True, False, 11, 12),
    ("Token-level Integrated Gradients attributions for a representative test-set patient. "
     "Tokens are ranked by absolute attribution magnitude. Red bars indicate risk-driving "
     "tokens (positive IG); blue bars indicate protective tokens (negative IG).",
     False, False, 11, 2),
    ("Alt text: Horizontal bar chart showing the top 20 salient tokens for one patient, "
     "colored red for positive attribution and blue for negative attribution.",
     False, True, 10, 2),

    ("Figure 3. Global attribution analysis.",
     True, False, 11, 12),
    ("Three-panel summary of population-level attribution on the test set. "
     "Panel A: population-level token importance (prevalence-weighted mean IG across all patients). "
     "Panel B: carrier-level conditional attribution (mean IG among patients who carry each token). "
     "Panel C: cross-visit temporal patterns ranked by summed attribution among carriers.",
     False, False, 11, 2),
    ("Alt text: Three-panel figure. Left panel shows a bar chart of population-level token "
     "importance. Middle panel shows carrier-level conditional attribution. Right panel shows "
     "ranked temporal patterns with their summed attribution scores.",
     False, True, 10, 2),

    ("Figure 4. Perturbation validation.",
     True, False, 11, 12),
    ("Dumbbell plot showing change in predicted 30-day readmission probability (Δŷ) after "
     "masking each of the top 15 temporal patterns from Panel C. Positive Δŷ indicates "
     "the masked tokens were protective; negative Δŷ indicates risk-driving tokens. "
     "Arrow color reflects direction of change: red for increase, blue for decrease.",
     False, False, 11, 2),
    ("Alt text: Dumbbell plot with 15 rows, one per pattern. Each row shows the baseline "
     "and masked predicted readmission probability connected by a colored arrow indicating "
     "the direction and magnitude of change after masking.",
     False, True, 10, 2),
]

for text, bold, italic, size, space_before in LEGENDS:
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)

print("    Figure Legends section appended")

# ── 9. Final scan for any remaining [val] placeholders ───────────────────────
print("\n[9] Final placeholder scan...")
IGNORE = {"[VISIT]", "[CLS]", "[PAD]", "[IG]", "[CPD]"}

def _hits(text):
    raw = re.findall(r"\[val\]|\[[A-Z_]{2,}\]", text)
    return [m for m in raw if m not in IGNORE]

remaining = []
for i, para in enumerate(doc.paragraphs):
    m = _hits(para.text)
    if m:
        remaining.append((i, para.text[:100], m))
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                m = _hits(para.text)
                if m:
                    remaining.append(("tbl", para.text[:100], m))

if remaining:
    print(f"  ⚠️  {len(remaining)} remaining:")
    for loc, snippet, matches in remaining:
        print(f"    [{loc}] {matches}  |  '{snippet}'")
else:
    print("  ✅ No placeholders remaining")

# ── 10. Abstract word count ──────────────────────────────────────────────────
# JAMIA structured abstracts have "Objective:", "Materials and Methods:", etc.
# as inline labels WITHIN the abstract. Stop at the first major top-level
# section (Background and Significance / Introduction / Keywords).
print("\n[10] Abstract word count...")
MAJOR = ("background and significance", "introduction", "keywords")
in_abstract, count = False, 0
for para in doc.paragraphs:
    t = para.text.strip()
    if not in_abstract and t.lower() == "abstract":
        in_abstract = True
        continue
    if in_abstract:
        if any(t.lower().startswith(k) for k in MAJOR):
            break
        count += len(t.split())
print(f"  ~{count} words {'✅' if count <= 250 else f'⚠️  over by {count-250}'}")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"\n✅ Saved: {DST}")
